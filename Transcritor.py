import os
import sys
import threading
import json
import whisper
import customtkinter as ctk
from tkinter import filedialog, messagebox
from datetime import datetime

from docx import Document
from fpdf import FPDF


def configurar_ffmpeg():
    if getattr(sys, 'frozen', False):
        pasta_base = sys._MEIPASS
    else:
        pasta_base = os.path.dirname(os.path.abspath(__file__))
    os.environ["PATH"] += os.pathsep + pasta_base


configurar_ffmpeg()

ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")


class TranscritorApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Transcritor de Mídia")
        self.geometry("500x500")
        self.caminho_video = None
        self.cancelado = False

        self.lbl_titulo = ctk.CTkLabel(self, text="Transcritor Whisper", font=("Arial", 20, "bold"))
        self.lbl_titulo.pack(pady=20)

        self.btn_selecionar = ctk.CTkButton(self, text="Selecionar Arquivo", command=self.selecionar_midia)
        self.btn_selecionar.pack(pady=10)

        self.lbl_arquivo = ctk.CTkLabel(self, text="Nenhum arquivo selecionado", text_color="gray")
        self.lbl_arquivo.pack(pady=5)

        self.lbl_formato = ctk.CTkLabel(self, text="Formato de saída:", font=("Arial", 12))
        self.lbl_formato.pack(pady=(15, 2))

        self.seletor_formato = ctk.CTkOptionMenu(
            self,
            values=["Markdown (.md)", "Texto (.txt)", "JSON (.json)", "Word (.docx)", "PDF (.pdf)"]
        )
        self.seletor_formato.pack(pady=5)

        self.btn_transcrever = ctk.CTkButton(self, text="Iniciar Transcrição", command=self.iniciar_transcricao,
                                             state="disabled")
        self.btn_transcrever.pack(pady=15)

        # Aqui entra o botão de cancelamento
        self.btn_cancelar = ctk.CTkButton(self, text="Cancelar Transcrição", command=self.cancelar_transcricao,
                                          fg_color="gray", state="disabled")
        self.btn_cancelar.pack(pady=5)

        self.barra_progresso = ctk.CTkProgressBar(self, mode="indeterminate")

        self.lbl_status = ctk.CTkLabel(self, text="")
        self.lbl_status.pack(pady=10)

    def selecionar_midia(self):
        caminho = filedialog.askopenfilename(
            title="Selecione a mídia",
            filetypes=[("Mídia", "*.mp4 *.mkv *.avi *.mp3 *.wav"), ("Todos os arquivos", "*.*")]
        )
        if caminho:
            self.caminho_video = caminho
            self.lbl_arquivo.configure(text=os.path.basename(self.caminho_video))
            self.btn_transcrever.configure(state="normal")
            self.lbl_status.configure(text="Pronto para começar.")

    def iniciar_transcricao(self):
        escolha = self.seletor_formato.get()
        if "Texto" in escolha:
            extensao = ".txt"
        elif "JSON" in escolha:
            extensao = ".json"
        elif "Word" in escolha:
            extensao = ".docx"
        elif "PDF" in escolha:
            extensao = ".pdf"
        else:
            extensao = ".md"

        nome_base, _ = os.path.splitext(self.caminho_video)
        self.nome_final = f"{nome_base}{extensao}"

        # Checagem acontece aqui, antes de travar a tela ou processar o áudio
        if os.path.exists(self.nome_final):
            sobrescrever = messagebox.askyesno(
                "Aviso de Substituição",
                f"O arquivo '{os.path.basename(self.nome_final)}' já existe.\nDeseja apagá-lo e criar um novo?"
            )
            if not sobrescrever:
                return

        self.cancelado = False
        self.btn_selecionar.configure(state="disabled")
        self.btn_transcrever.configure(state="disabled")
        self.seletor_formato.configure(state="disabled")
        self.btn_cancelar.configure(state="normal", fg_color="#D32F2F")

        self.lbl_status.configure(text="Processando... Isso pode demorar.", text_color="yellow")
        self.barra_progresso.pack(pady=5)
        self.barra_progresso.start()

        threading.Thread(target=self.processar_video, daemon=True).start()

    def cancelar_transcricao(self):
        self.cancelado = True
        self.lbl_status.configure(text="Cancelando... Aguarde o processo fechar.", text_color="orange")
        self.resetar_interface()

    def resetar_interface(self):
        self.barra_progresso.stop()
        self.barra_progresso.pack_forget()
        self.btn_selecionar.configure(state="normal")
        self.btn_transcrever.configure(state="normal")
        self.seletor_formato.configure(state="normal")
        self.btn_cancelar.configure(state="disabled", fg_color="gray")

    def processar_arquivo(self):
        try:
            inicio = datetime.now()
            modelo = whisper.load_model("small")
            resultado = modelo.transcribe(self.caminho_video)

            if self.cancelado:
                self.lbl_status.configure(text="Transcrição abortada.", text_color="orange")
                return

            texto_gerado = resultado["text"]
            escolha = self.seletor_formato.get()

            if "JSON" in escolha:
                with open(self.nome_final, "w", encoding="utf-8") as arquivo:
                    json.dump(resultado, arquivo, indent=4, ensure_ascii=False)
            elif "Word" in escolha:
                doc = Document()
                doc.add_heading('Transcrição', 0)
                doc.add_paragraph(texto_gerado)
                doc.save(self.nome_final)
            elif "PDF" in escolha:
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("helvetica", size=12)
                pdf.multi_cell(0, 10, txt=texto_gerado)
                pdf.output(self.nome_final)
            else:
                with open(self.nome_final, "w", encoding="utf-8") as arquivo:
                    arquivo.write(texto_gerado)

            if not self.cancelado:
                duracao = datetime.now() - inicio
                tempo_formatado = str(duracao).split('.')[0]
                self.lbl_status.configure(text=f"Sucesso! Arquivo salvo.\nTempo gasto: {tempo_formatado}",
                                          text_color="green")

        except Exception as e:
            if not self.cancelado:
                self.lbl_status.configure(text=f"Erro:\n{str(e)}", text_color="red")

        finally:
            if not self.cancelado:
                self.resetar_interface()


if __name__ == "__main__":
    app = TranscritorApp()
    app.mainloop()